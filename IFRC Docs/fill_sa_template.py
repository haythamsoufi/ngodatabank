"""
One-off script: fill IFRC SA Simplified Template v4 2026 from PCH / UR / KT docs.
Requires: pip install python-docx (available in Backoffice venv).
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.table import Table


def set_paragraph_text(paragraph, text: str) -> None:
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def set_cell_text(cell, text: str) -> None:
    for p in cell.paragraphs:
        set_paragraph_text(p, "")
    if cell.paragraphs:
        set_paragraph_text(cell.paragraphs[0], text)
    else:
        cell.add_paragraph(text)


def replace_para_by_prefix(doc: Document, prefix: str, new_line: str) -> bool:
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith(prefix):
            set_paragraph_text(p, new_line)
            return True
    return False


def fill_field_table(table: Table, fills: dict[str, str]) -> None:
    for row in table.rows[1:]:
        key = row.cells[0].text.strip()
        if key in fills:
            set_cell_text(row.cells[1], fills[key])


def main() -> None:
    base = Path(__file__).resolve().parent
    src = base / "SA_Simplified_ Template _v4 2026.docx"
    backup = base / "SA_Simplified_ Template _v4 2026.template_backup.docx"
    if not backup.exists():
        shutil.copy2(src, backup)

    doc = Document(str(src))

    # Title line — mark as programme-filled draft
    set_paragraph_text(doc.paragraphs[0], "Solution Architecture Document (Draft — IFRC Network Data Portal)")

    # Application name (blank lines after label)
    set_paragraph_text(doc.paragraphs[5], "IFRC Network Data Portal")

    # Document authorization table (table 0)
    auth = doc.tables[0]
    set_cell_text(auth.rows[1].cells[1], "Haytham Al Soufi")
    set_cell_text(auth.rows[1].cells[2], "Senior Officer, Federation-Wide Data Analysis (primary author)")
    set_cell_text(auth.rows[2].cells[1], "Rania Alereksoussi")
    set_cell_text(auth.rows[2].cells[2], "Manager, Federation-Wide RBM & Data Systems (business owner)")
    set_cell_text(auth.rows[3].cells[1], "Marie-France Shakour (PM); Carlota Tarazona (PO); DTD/Infra/Dev teams per PCH")
    set_cell_text(auth.rows[3].cells[2], "Project steering, product ownership, EA/infra/application support")
    set_cell_text(auth.rows[4].cells[1], "TBD per DTD governance")
    set_cell_text(auth.rows[4].cells[2], "Formal sign-off when review cycle completes")

    # --- Section 1 ---
    replace_para_by_prefix(
        doc,
        "In Scope:",
        "In Scope: Unified federation-wide structured data collection and management: form/template design "
        "(Form Builder), assignments to countries/NS and entities, Indicator Bank–linked indicators, focal-point "
        "data entry with validation and audit trail, exports (e.g. Excel), REST APIs (including mobile JWT surface), "
        "notifications (in-app/email/push where configured), multilingual content, analytics dashboards, and phased "
        "AI/RAG document features. MVP alignment: Unified Mid-Year Reporting 2026 and related PMER/FDS processes "
        "(FDRS, Unified Planning/Reporting, fed-wide emergencies) per PCH.",
    )
    replace_para_by_prefix(
        doc,
        "Out of Scope:",
        "Out of Scope: Full one-shot migration of all historical legacy systems and datasets at initial launch; "
        "complete sunsetting of Kobo/Excel/Power Platform in one release; full public portal publishing scope unless "
        "explicitly required for the reporting round; some advanced analytics/AI integrations may evolve post go-live "
        "per PCH and knowledge-transfer notes.",
    )
    replace_para_by_prefix(
        doc,
        "Primary Readers:",
        "Primary Readers: DTD enterprise & solution architects; infrastructure and security teams; application support; "
        "project manager and PMO; product owner and FDS technical administrators.",
    )
    replace_para_by_prefix(
        doc,
        "Secondary Readers:",
        "Secondary Readers: Business owner and sponsors; National Society and IFRC regional focal points (operational context); "
        "senior management consuming federation-wide reporting.",
    )

    # --- Section 2 ---
    replace_para_by_prefix(
        doc,
        "Project / CR Number:",
        "Project / CR Number: As recorded in IFRC project register (PCH: IFRC Network Data Portal).",
    )
    replace_para_by_prefix(
        doc,
        "Solution Name:",
        "Solution Name: IFRC Network Data Portal (Backoffice web application; monorepo also contains Website/MobileApp for future use).",
    )
    replace_para_by_prefix(
        doc,
        "Solution Type:",
        "Solution Type: Custom-built web application (IFRC-owned codebase; not commercial SaaS).",
    )
    replace_para_by_prefix(
        doc,
        "Vendor / Product:",
        "Vendor / Product: IFRC in-house delivery (AI-assisted development). Runtime dependencies include Microsoft Azure "
        "(App Service, Container Registry, Blob Storage, Entra ID B2C), PostgreSQL, and optional third-party AI APIs "
        "(OpenAI / Azure OpenAI / Gemini) when AI features are enabled.",
    )
    replace_para_by_prefix(
        doc,
        "Expected Go-live Date:",
        "Expected Go-live Date: April–May 2026 — production rollout targeted before Unified Mid-Year Reporting data collection.",
    )
    replace_para_by_prefix(
        doc,
        "Project Manager:",
        "Project Manager: Marie-France Shakour (DTD Project Manager).",
    )

    replace_para_by_prefix(
        doc,
        "Business Need:",
        "Business Need: Replace fragmented tools (Kobo, Excel, ad hoc forms, multiple platforms) with one standardized "
        "platform that enforces Indicator Bank definitions, reduces duplicate NS reporting, improves data quality, and "
        "enables analytics and controlled reuse across the IFRC network.",
    )
    replace_para_by_prefix(
        doc,
        "Current State:",
        "Current State: Parallel systems and manual consolidation; inconsistent KPI names/definitions; higher maintenance "
        "cost; data loss/reuse limits; uneven access across teams (per PCH and knowledge-transfer documentation).",
    )
    replace_para_by_prefix(
        doc,
        "Desired Outcome:",
        "Desired Outcome: Single entry point for configured federation-wide processes; standardized indicators; "
        "federation-wide accessibility; faster reporting cycle support; reduced operational tool sprawl; auditable "
        "admin and user actions; APIs/exports for downstream consumers.",
    )

    replace_para_by_prefix(
        doc,
        "Risk 1:",
        "Risk 1: Delay or failure to launch before Unified Mid-Year Reporting — NS/IFRC teams fall back to fragmented tools, "
        "duplication continues, and Indicator Bank standardisation is harder to enforce.",
    )
    replace_para_by_prefix(
        doc,
        "Risk 2:",
        "Risk 2: Security or B2C authentication issues block trusted rollout — reputational risk, login confusion, or blocked "
        "access during live reporting windows.",
    )
    replace_para_by_prefix(
        doc,
        "Risk 3:",
        "Risk 3: Insufficient support/runbook clarity post-handover — slower incident resolution and higher operational "
        "disruption during critical collection periods.",
    )

    replace_para_by_prefix(
        doc,
        "Connected Systems:",
        "Connected Systems: Azure AD B2C (OIDC SSO); Azure Blob Storage (files); IFRC Email API; IFRC Translation API / "
        "LibreTranslate; optional KoBo import; optional AI providers; PostgreSQL; public/website/mobile consumers via REST "
        "where enabled.",
    )
    replace_para_by_prefix(
        doc,
        "Dependencies:",
        "Dependencies: IFRC Azure landing zone (App Service, ACR, networking); DTD code review and penetration testing "
        "closure; Indicator Bank content and governance; NS/PMER participation in UAT; valid TLS and secrets management.",
    )
    replace_para_by_prefix(
        doc,
        "Impacted Teams:",
        "Impacted Teams: FDS/PMER (configuration and coordination); NS focal points (data entry); IFRC regional/country offices; "
        "DTD infrastructure, application support, security; EA/PMO; strategic planning consumers of outputs.",
    )

    # --- Section 3 ---
    proc = (
        "Key processes supported: (1) Template design & versioning — administrators define sections/items, indicators, "
        "validation and skip logic. (2) Assignment management — assign forms to countries/entities for reporting periods; "
        "track status. (3) Data capture — focal points enter/submit with auto-save, presence awareness, attachments where "
        "allowed. (4) Review/coordination — coordinators monitor completeness and drive corrections (per user requirements). "
        "(5) Publication of data for authorised consumers — exports, APIs, dashboards. (6) Operational notifications and "
        "reminders during cycles."
    )
    for p in doc.paragraphs:
        if p.text.strip().startswith("Describe the key business processes"):
            set_paragraph_text(p, proc)
            break

    use_cases = (
        "System administrators: build templates, manage Indicator Bank links, assignments, users/RBAC, translations, "
        "resources. Data contributors (NS focal points / IFRC delegates): complete assigned forms, upload evidence. "
        "Coordinators/validators: monitor progress, validate data quality, coordinate clarifications. Data consumers: "
        "export or query via APIs/reports. DT operators: deploy, monitor, respond to incidents."
    )
    for p in doc.paragraphs:
        if p.text.strip().startswith("List who will use the system"):
            set_paragraph_text(p, use_cases)
            break

    arch_text = (
        "High-level structure: Browser clients hit Azure App Service hosting a Python Flask application (Gunicorn). "
        "Server-rendered Jinja2 UI (TailwindCSS) for admin and form entry; REST API v1 for integrations; optional WebSockets "
        "for notifications/AI where enabled. Core modules: authentication (local + B2C OIDC), modular admin (form builder, "
        "assignments, content, system admin, analytics, utilities), forms entry, API layer, AI (HTTP/SSE/WS), notifications. "
        "PostgreSQL 15 (pgvector for embeddings when RAG enabled) stores relational and JSON form payloads. Azure Blob "
        "stores uploads. CI/CD via GitHub Actions (tests, CodeQL, container build/push, deploy workflow to App Service)."
    )
    for p in doc.paragraphs:
        if p.text.strip().startswith("Provide a high-level diagram"):
            set_paragraph_text(p, arch_text)
            break

    hint411 = (
        "Layers: (B) Federation-wide reporting & indicator governance. (I) Form templates, submissions, Indicator Bank "
        "metadata, audit logs, documents. (A) Flask app + APIs + optional Next.js public site / Flutter mobile consumers. "
        "(T) Azure App Service Linux containers, ACR, PostgreSQL, Blob, B2C. (S) RBAC, CSRF, CSP, rate limits, encryption "
        "in transit/at rest, security logging."
    )
    for p in doc.paragraphs:
        if p.text.strip().startswith("Try and represent the layers"):
            set_paragraph_text(p, hint411)
            break

    # BIATS table — column "Description | Example" is index 2 on 3-col table
    biats = doc.tables[1]
    examples = [
        (
            "B",
            "Federation-wide data collection cycles (e.g. Unified Mid-Year Reporting): template configuration, NS/IFRC "
            "contributions, validation/coordination, and management reporting.",
        ),
        (
            "I",
            "Structured submission payloads (JSON per form/section), Indicator Bank definitions, multilingual labels, "
            "assignment status, audit/admin logs, document metadata, optional AI embeddings.",
        ),
        (
            "A",
            "Flask monolith + blueprints; Jinja/Tailwind UI; REST API; optional AI and notification WebSockets; integrations "
            "listed under APIs & Integrations.",
        ),
        (
            "T",
            "Azure App Service (staging/production), Docker images via Azure Container Registry, PostgreSQL 15, Azure Blob, "
            "Entra External ID (B2C); GitHub Actions pipelines.",
        ),
        (
            "S",
            "B2C + local auth, RBAC/permissions, MFA via IdP where enforced, TLS, secure cookies, CSRF, CSP nonces, "
            "rate limiting, admin/user/security event logging, penetration testing and dependency scanning in CI.",
        ),
    ]
    for letter, desc in examples:
        for row in biats.rows[1:]:
            if row.cells[0].text.strip() == letter:
                set_cell_text(row.cells[2], desc)
                break

    # Software components table — replace example row and append rows
    comp = doc.tables[2]
    specs = [
        ("Web UI (admin & forms)", "Jinja2 + TailwindCSS + JS", "Staff/NS browser experience for builder, entry, dashboards"),
        ("API layer", "Flask REST (`/api/v1`, mobile `/api/mobile/v1`)", "Integrations, mobile app, automation consumers"),
        ("Application core", "Python 3.11, Flask 3.x, SQLAlchemy", "Business logic, RBAC, workflows, validation"),
        ("Data tier", "PostgreSQL 15 + pgvector", "Transactional storage, JSON form data, optional vector search"),
        ("Auth", "Flask-Login, B2C OIDC, JWT (mobile)", "Interactive sessions and bearer access for APIs"),
        ("Files", "Azure Blob (+ local fallback)", "Uploads, resources, generated exports/thumbnails"),
        ("Observability", "App logs, Azure Monitor-ready", "Health, diagnostics, security events"),
    ]
    # Clear all but header
    while len(comp.rows) > 1:
        tr = comp.rows[-1]._tr
        tr.getparent().remove(tr)
    for name, tech, purpose in specs:
        row = comp.add_row()
        set_cell_text(row.cells[0], name)
        set_cell_text(row.cells[1], tech)
        set_cell_text(row.cells[2], purpose)

    # APIs table (table 3)
    api_t = doc.tables[3]
    set_cell_text(api_t.rows[1].cells[1], "IFRC Website/portal (when used), internal tools, approved automation, Flutter mobile app via mobile API, analysts/scripts with JWT/API keys where configured.")
    set_cell_text(api_t.rows[2].cells[1], "Azure AD B2C; Azure Blob; IFRC Email API; IFRC Translation API / LibreTranslate; optional OpenAI/Azure OpenAI/Gemini; KoBo (import); FCM (push) when enabled.")
    set_cell_text(api_t.rows[3].cells[1], "Primarily HTTPS REST/JSON; OIDC/OAuth2 for SSO; WebSockets for select features; email webhooks/API.")

    # 5.1 Data overview
    fill_field_table(
        doc.tables[4],
        {
            "Data Types": "Structured indicator/question responses; organizational metadata; user directory data (names, emails); optional documents; operational logs; optional AI conversation/document chunks. Mix of Internal operational data; some indicators/public outputs may be non-personal or public-facing when published elsewhere.",
            "Data Sources": "Human entry via web UI; bulk Excel import/export; KoBo imports; administrative configuration; Indicator Bank master data; translation services; optional external AI inference.",
            "Data Destinations": "Primary store: Azure-hosted PostgreSQL; files in Azure Blob; exports downloaded by users; downstream analytics via API/export (Power BI or other tools out of band unless integrated later).",
            "Data Retention": "To be formalised in records-management policy; currently operational retention with audit logs and backups per infrastructure settings (TBD with DTD — align to IFRC retention classes).",
            "Reporting / Analytics": "In-app analytics/admin dashboards; Excel exports; API consumption; future/optional self-service BI not required for MVP per draft UR scope note.",
        },
    )

    fill_field_table(
        doc.tables[5],
        {
            "Reporting Tool": "Built-in admin/analytics views; Excel exports; optional external BI (e.g. Power BI) fed via export/API.",
            "Key Reports": "Assignment progress, submission status, user activity/session reporting, audit trail views, indicator datasets for reporting cycles.",
            "Data Platform": "Not yet stated — interface to enterprise data platform TBD with EA/data architecture; APIs and exports provided as integration points.",
            "Data Consumers": "FDS/PMER, regional desks, strategy/PRD analysts, NS focal points (own data), authorised HQ roles.",
            "Refresh Frequency": "Near real-time for operational UI; batch/ondemand for heavy exports; analytics aligned to reporting windows.",
        },
    )

    fill_field_table(
        doc.tables[6],
        {
            "Hosting Type": "Microsoft Azure (PaaS App Service).",
            "Server Count": "Managed App Service instances (staging + production); exact SKU/count per Azure subscription design.",
            "Operating System": "Linux containers (Python slim base image) on App Service.",
            "CPU / RAM / Storage": "Per App Service plan sizing; PostgreSQL tier for workload; Blob for file volume — formal sizing with infra team.",
            "Network / Firewall Rules": "Standard IFRC Azure network controls; TLS public endpoints; restricted admin access; secrets in App Service configuration/Key Vault pattern per org standard.",
        },
    )

    fill_field_table(
        doc.tables[7],
        {
            "Database Type": "PostgreSQL 15 with pgvector extension (for optional embeddings/RAG).",
            "Hosting": "Azure Database for PostgreSQL or equivalent managed service used by the project environment.",
            "Estimated Size": "Grows with reporting cycles, attachments, and audit volume; capacity planning with DTD infra.",
            "Backup Frequency": "Per Azure managed database policy (typically automated daily with PITR where enabled).",
            "Retention Period": "Per IFRC/DTD backup retention standard for managed DBs (confirm with infra).",
        },
    )

    env = doc.tables[8]
    set_cell_text(env.rows[1].cells[1], "Developer workstations / local Docker compose for feature work.")
    set_cell_text(env.rows[1].cells[2], "Local README / CLAUDE.md quickstart.")
    set_cell_text(env.rows[2].cells[1], "Staging App Service — integration, UAT, pen-test fixes, release candidates.")
    set_cell_text(env.rows[2].cells[2], "Azure staging slot / staging web app URL (per deployment config).")
    set_cell_text(env.rows[3].cells[1], "Live federation data collection and admin operations.")
    set_cell_text(env.rows[3].cells[2], "Azure production App Service endpoint (internal DNS/public as configured).")

    # Personal data & availability & IAM
    fill_field_table(
        doc.tables[11],
        {
            "Contains Personal Data?": "Yes — user accounts and operational contacts.",
            "Type of Personal Data": "Names, email addresses, roles, login metadata; optional profile fields; activity timestamps.",
            "Data Subject": "IFRC staff, consultants, and authorised National Society users using the platform.",
            "Privacy Impact Assessment": "Confirm with DPO/DTD — required where personal data processing; document outcome when completed.",
        },
    )

    fill_field_table(
        doc.tables[12],
        {
            "Required Availability": "High during active reporting windows; target aligned with App Service SLA and business expectations (define explicit % with stakeholders).",
            "Recovery Time Objective (RTO)": "To be agreed with infra (e.g. hours-to-1 business day for major outage; tune with criticality).",
            "Recovery Point Objective (RPO)": "Aligned to managed DB backup/PITR capabilities (typically low minutes to sub-hour where PITR enabled).",
            "Criticality": "High for scheduled global reporting campaigns; medium during inter-cycle periods.",
        },
    )

    fill_field_table(
        doc.tables[13],
        {
            "Authentication Method": "Azure AD B2C (OIDC) for IFRC SSO track; local email/password fallback where configured; mobile JWT issuance for mobile API.",
            "Single Sign-On (SSO)": "Yes (B2C) for the primary staff path where enabled.",
            "User Roles": "RBAC-backed roles (e.g. system admin, focal_point, view_only) plus granular permissions; entity/country scoping for assignments.",
            "Multi-Factor Authentication": "As enforced by organisational IdP policy for B2C users (Yes where MFA mandated by tenant).",
            "Auditing / Logging": "Yes — admin action logs, meaningful activity events, security events, session tracking patterns per knowledge-transfer doc.",
        },
    )

    fill_field_table(
        doc.tables[14],
        {
            "Support Team": "DTD Application Support / Digital Unit handover model; FDS product owner for functional triage; escalation per SteerCo agreement.",
            "Support Hours": "Geneva business hours baseline; critical incident handling per IFRC major incident process.",
            "Escalation Path": "Service Desk → Application support → Infra/security → vendor (Azure) via ITIL channels as defined by DTD.",
            "Service Desk": "To be registered in MSM / service catalogue as part of go-live handover (per template question).",
        },
    )

    fill_field_table(
        doc.tables[15],
        {
            "Software / Service": "Azure App Service, Azure Database for PostgreSQL, Azure Blob, Entra External ID (B2C); optional OpenAI/Azure OpenAI; GitHub Enterprise usage.",
            "License Type": "Organisational subscriptions / PAYG cloud; API usage-based for AI.",
            "Number of Licenses": "Not per-seat for core cloud; AI tokens per provider contract.",
            "Annual Cost": "Per IFRC cloud chargeback / subscription records (fill from finance).",
            "Renewal Date": "Per Azure and vendor contract schedule.",
        },
    )

    fill_field_table(
        doc.tables[16],
        {
            "Monitoring Tool": "Azure Monitor / App Service diagnostics; application logs; optional Application Insights.",
            "What is Monitored": "HTTP health, errors, latency, dependency failures, container restarts, DB connectivity, auth anomalies.",
            "Alert Recipients": "On-call distribution list per DTD infra; application support for app-level alerts.",
            "Escalation": "Sev-based routing per DTD operational playbook.",
        },
    )

    fill_field_table(
        doc.tables[17],
        {
            "Functional Testing": "pytest suite in CI (~50% coverage gate); route and service tests for forms, API, auth, notifications.",
            "UAT Results": "Link to UAT sign-off / test evidence folder when available (PMER/NS pilot).",
            "Technical Testing": "CodeQL, dependency review, Bandit/Gitleaks in CI; penetration testing March 2026 per PCH; performance tuning for large forms/exports.",
        },
    )

    fill_field_table(
        doc.tables[18],
        {
            "Initial Training": "Regional online demos for admins and focal points; onboarding per PCH training plan (Q1–Q2 2026).",
            "Ongoing / Recurrent Training": "New focal point onboarding; release notes when major features ship.",
            "Training Materials": "Backoffice `docs/` multilingual guides; knowledge-transfer document; future LMS links TBD.",
        },
    )

    fill_field_table(
        doc.tables[19],
        {
            "DR Procedure": "Redeploy last known-good container from ACR; restore PostgreSQL from backup/PITR; revalidate secrets and B2C config; communications template per DTD DR.",
            "RTO": "Match table 12 — confirm numerically with infra.",
            "RPO": "Match table 12 — confirm with DB backup design.",
            "Last DR Test": "Planned / TBD — schedule with infra after go-live.",
        },
    )

    rel = doc.tables[20]
    set_cell_text(rel.rows[1].cells[0], "PCH-001")
    set_cell_text(rel.rows[1].cells[1], "Project Charter — IFRC Network Data Portal")
    set_cell_text(rel.rows[1].cells[2], "As approved / in circulation")
    # Add second related doc row
    rel.add_row()
    set_cell_text(rel.rows[2].cells[0], "UR-001")
    set_cell_text(rel.rows[2].cells[1], "Draft User Requirements — IFRC Network Data Portal")
    set_cell_text(rel.rows[2].cells[2], "Draft")
    rel.add_row()
    set_cell_text(rel.rows[3].cells[0], "KT-001")
    set_cell_text(rel.rows[3].cells[1], "Application Development & Knowledge Transfer — IFRC Network Data Portal")
    set_cell_text(rel.rows[3].cells[2], "1.0 / Apr 2026")

    rev = doc.tables[21]
    set_cell_text(rev.rows[1].cells[0], "0.1")
    set_cell_text(rev.rows[1].cells[1], "April 2026")
    set_cell_text(rev.rows[1].cells[2], "Auto-filled draft from PCH, draft UR, and knowledge-transfer documentation.")

    replace_para_by_prefix(
        doc,
        "Input:",
        "Input: Administrators configure templates and assignments; contributors enter structured responses, uploads, and "
        "imports (Excel/KoBo where used); integrations post/pull via APIs; IdP issues tokens for SSO users.",
    )
    replace_para_by_prefix(
        doc,
        "Processing:",
        "Processing: Validation rules and skip logic; multilingual rendering; RBAC enforcement; versioning; optional "
        "translation and AI enrichment; notifications triggered on state changes.",
    )
    replace_para_by_prefix(
        doc,
        "Output / Storage:",
        "Output / Storage: Authoritative data persisted in PostgreSQL; artefacts in Blob; users download exports; "
        "authorised consumers read JSON via REST; analytics aggregates in application views.",
    )
    replace_para_by_prefix(
        doc,
        "External Sharing:",
        "External Sharing: Cloud AI inference when enabled (content minimisation policies apply); email via IFRC Email API; "
        "no unmanaged public dump of restricted datasets — publishing to public channels only through governed routes.",
    )

    doc.save(str(src))
    print("Saved:", src)
    print("Backup:", backup if backup.exists() else "n/a")


if __name__ == "__main__":
    main()
