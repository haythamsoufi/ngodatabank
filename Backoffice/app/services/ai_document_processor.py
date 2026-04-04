"""
AI Document Processor Service

Handles loading and processing of multi-format documents for RAG system.
Supports PDF, Word, Excel, Text, and Markdown files.
"""

import os
import hashlib
import logging
from typing import Dict, List, Optional, Tuple, Any
from pathlib import Path
import mimetypes
import re

from flask import current_app

logger = logging.getLogger(__name__)

_YEAR_RE = re.compile(r"\b(19\d{2}|20\d{2})\b")
_AMOUNT_RE = re.compile(r"\b\d+(?:\.\d+)?\s*[MK]\b|\b\d{1,3}(?:,\d{3})+\b", re.IGNORECASE)
_PDF_CORRUPTION_HINTS = (
    "object out of range",
    "xref size",
    "xref",
    "expected object number",
    "cannot open broken document",
    "format error",
    "syntax error",
    "cannot find startxref",
    "repairing",
)


class DocumentProcessingError(Exception):
    """Raised when document processing fails."""
    pass


class AIDocumentProcessor:
    """
    Service for processing documents into text and metadata.

    Supports multiple document formats and extracts:
    - Text content with layout preservation
    - Metadata (title, author, creation date, etc.)
    - Images (for multi-modal processing)
    - Page numbers and sections
    """

    # Supported file types
    SUPPORTED_TYPES = {
        'pdf': ['.pdf'],
        'word': ['.docx', '.doc'],
        'excel': ['.xlsx', '.xls', '.csv'],
        'text': ['.txt', '.text'],
        'markdown': ['.md', '.markdown'],
        'html': ['.html', '.htm'],
    }

    def __init__(self):
        """Initialize the document processor."""
        self.max_file_size = int(current_app.config.get('AI_MAX_DOCUMENT_SIZE_MB', 50)) * 1024 * 1024  # Convert to bytes

    @staticmethod
    def _is_probably_corrupt_pdf_error(err: Exception) -> bool:
        txt = str(err or "").lower()
        return any(h in txt for h in _PDF_CORRUPTION_HINTS)

    def is_supported_file(self, filename: str) -> bool:
        """Check if file type is supported."""
        ext = Path(filename).suffix.lower()
        return any(ext in extensions for extensions in self.SUPPORTED_TYPES.values())

    def get_file_type(self, filename: str) -> Optional[str]:
        """Get the file type category."""
        ext = Path(filename).suffix.lower()
        for file_type, extensions in self.SUPPORTED_TYPES.items():
            if ext in extensions:
                return file_type
        return None

    def calculate_content_hash(self, file_path: str) -> str:
        """Calculate SHA256 hash of file content for deduplication."""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()

    def process_document(
        self,
        file_path: str,
        filename: str,
        extract_images: bool = False,
        ocr_enabled: bool = False,
        max_pages: Optional[int] = None,
    ) -> Dict[str, Any]:
        """
        Process a document and extract text, metadata, and optionally images.

        Args:
            file_path: Path to the document file
            filename: Original filename
            extract_images: Whether to extract images (for multi-modal)
            ocr_enabled: Whether to use OCR for scanned PDFs

        Returns:
            Dictionary containing:
                - text: Extracted text content
                - metadata: Document metadata (title, author, pages, etc.)
                - images: List of extracted images (if extract_images=True)
                - sections: List of document sections with page numbers

        Raises:
            DocumentProcessingError: If processing fails
        """
        # Validate file exists
        if not os.path.exists(file_path):
            raise DocumentProcessingError(f"File not found: {file_path}")

        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size > self.max_file_size:
            raise DocumentProcessingError(
                f"File too large: {file_size / 1024 / 1024:.1f}MB "
                f"(max: {self.max_file_size / 1024 / 1024:.1f}MB)"
            )

        # Get file type
        file_type = self.get_file_type(filename)
        if not file_type:
            raise DocumentProcessingError(f"Unsupported file type: {filename}")

        logger.info(f"Processing {file_type} document: {filename} ({file_size / 1024:.1f}KB)")

        # Route to appropriate processor
        try:
            if file_type == 'pdf':
                return self._process_pdf(file_path, filename, extract_images, ocr_enabled, max_pages=max_pages)
            elif file_type == 'word':
                return self._process_word(file_path, filename)
            elif file_type == 'excel':
                return self._process_excel(file_path, filename)
            elif file_type == 'text':
                return self._process_text(file_path, filename)
            elif file_type == 'markdown':
                return self._process_markdown(file_path, filename)
            elif file_type == 'html':
                return self._process_html(file_path, filename)
            else:
                raise DocumentProcessingError(f"No processor for file type: {file_type}")
        except Exception as e:
            logger.error(f"Error processing {filename}: {str(e)}", exc_info=True)
            raise DocumentProcessingError("Failed to process document.")

    def _process_pdf(
        self,
        file_path: str,
        filename: str,
        extract_images: bool,
        ocr_enabled: bool,
        *,
        max_pages: Optional[int] = None,
    ) -> Dict[str, Any]:
        """Process PDF document."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise DocumentProcessingError("PyMuPDF (fitz) not installed. Run: pip install PyMuPDF")

        # Keep MuPDF parser warnings from flooding logs/stdout by default.
        # Enable explicit low-level warnings when troubleshooting parser issues.
        try:
            verbose_mupdf = bool(current_app.config.get("AI_PDF_MUPDF_VERBOSE_ERRORS", False))
        except Exception as e:
            logger.debug("AI_PDF_MUPDF_VERBOSE_ERRORS parse failed: %s", e)
            verbose_mupdf = False
        if not verbose_mupdf:
            try:
                tools = getattr(fitz, "TOOLS", None)
                if tools is not None:
                    if hasattr(tools, "mupdf_display_errors"):
                        tools.mupdf_display_errors(False)
                    if hasattr(tools, "mupdf_display_warnings"):
                        tools.mupdf_display_warnings(False)
            except Exception as e:
                logger.debug("mupdf_display_errors/warnings failed: %s", e)

        result = {
            'text': '',
            'metadata': {},
            'images': [],
            'sections': [],
            'tables': [],
            'pages': []
        }

        try:
            doc = fitz.open(file_path)

            # Extract metadata
            result['metadata'] = {
                'title': doc.metadata.get('title', filename),
                'author': doc.metadata.get('author', ''),
                'subject': doc.metadata.get('subject', ''),
                'creator': doc.metadata.get('creator', ''),
                'producer': doc.metadata.get('producer', ''),
                'creation_date': doc.metadata.get('creationDate', ''),
                'modification_date': doc.metadata.get('modDate', ''),
                'total_pages': len(doc),
                'format': doc.metadata.get('format', 'PDF'),
            }

            # Extract text page by page
            for page_num, page in enumerate(doc, start=1):
                if isinstance(max_pages, int) and max_pages > 0 and int(page_num) > int(max_pages):
                    break
                # Optional: capture word-level layout for robust visual extraction (UPR KPIs, etc.)
                # This is intentionally disabled by default due to payload size.
                try:
                    layout_words_enabled = bool(current_app.config.get("AI_PDF_LAYOUT_WORDS_ENABLED", False))
                except Exception as e:
                    logger.debug("AI_PDF_LAYOUT_WORDS_ENABLED parse failed: %s", e)
                    layout_words_enabled = False
                try:
                    layout_words_max_pages = int(current_app.config.get("AI_PDF_LAYOUT_WORDS_MAX_PAGES", 3))
                except Exception as e:
                    logger.debug("AI_PDF_LAYOUT_WORDS_MAX_PAGES parse failed: %s", e)
                    layout_words_max_pages = 3

                # Optional: capture a cropped rendering of the top-of-page region for vision-based extraction.
                # Also disabled by default; only used during chunking (not stored in DB).
                try:
                    upr_vision_enabled = bool(current_app.config.get("AI_UPR_VISION_KPI_ENABLED", False))
                except Exception as e:
                    logger.debug("AI_UPR_VISION_KPI_ENABLED parse failed: %s", e)
                    upr_vision_enabled = False
                try:
                    upr_vision_max_pages = int(current_app.config.get("AI_UPR_VISION_MAX_PAGES", 1))
                except Exception as e:
                    logger.debug("AI_UPR_VISION_MAX_PAGES parse failed: %s", e)
                    upr_vision_max_pages = 1
                try:
                    upr_vision_dpi = int(current_app.config.get("AI_UPR_VISION_DPI", 160))
                except Exception as e:
                    logger.debug("AI_UPR_VISION_DPI parse failed: %s", e)
                    upr_vision_dpi = 160
                try:
                    upr_vision_clip_top_frac = float(current_app.config.get("AI_UPR_VISION_CLIP_TOP_FRAC", 0.42))
                except Exception as e:
                    logger.debug("AI_UPR_VISION_CLIP_TOP_FRAC parse failed: %s", e)
                    upr_vision_clip_top_frac = 0.42

                # Extract tables (best-effort) first so we can optionally exclude table regions from text.
                page_tables: List[Dict[str, Any]] = []
                if current_app.config.get('AI_TABLE_EXTRACTION_ENABLED', True):
                    try:
                        page_tables = self._extract_tables_from_page(page, page_num) or []
                        if page_tables:
                            result['tables'].extend(page_tables)
                    except Exception as e_tables:
                        logger.debug(f"Table extraction failed on page {page_num}: {e_tables}", exc_info=True)

                exclude_table_text = bool(current_app.config.get("AI_EXCLUDE_TABLE_TEXT_FROM_PDF_TEXT", True))
                table_bboxes = [
                    t.get("bbox")
                    for t in (page_tables or [])
                    if isinstance((t or {}).get("bbox"), (list, tuple)) and len((t or {}).get("bbox")) == 4
                ]

                # Extract page text, optionally excluding table areas to avoid duplicate noisy "semantic" chunks.
                if exclude_table_text and table_bboxes:
                    text = self._extract_text_excluding_bboxes(page, table_bboxes)
                else:
                    # Prefer sorted extraction for better reading order in tables/columns.
                    # Older PyMuPDF versions may not support sort=; fall back gracefully.
                    try:
                        text = page.get_text("text", sort=True)
                    except TypeError:
                        text = page.get_text("text")

                # Try OCR if text is minimal and OCR is enabled.
                # IMPORTANT: avoid OCR on pages with detected tables when table text exclusion is enabled,
                # because OCR would re-introduce the table text we intentionally removed.
                if ocr_enabled and len((text or "").strip()) < 50 and not (exclude_table_text and table_bboxes):
                    text = self._ocr_page(page, page_num)

                # Always keep a page record even when OCR/text is empty.
                # UPR plan pages can be infographic-heavy and produce near-empty text, but we may still
                # want geometry / word-layout / vision crops for deterministic extraction.
                page_obj: Dict[str, Any] = {
                    'page_number': page_num,
                    'text': text or "",
                    'char_count': len(text or "")
                }

                if (text or "").strip():
                    result['text'] += f"\n\n[Page {page_num}]\n{text}"

                # Attach page geometry (useful for layout-based extractors).
                try:
                    page_obj["page_width"] = float(page.rect.width)
                    page_obj["page_height"] = float(page.rect.height)
                except Exception as e:
                    logger.debug("Page geometry extraction failed: %s", e)

                # Attach word-level layout data (first N pages only).
                if layout_words_enabled and int(page_num) <= int(layout_words_max_pages):
                    try:
                        # words(): (x0, y0, x1, y1, "word", block_no, line_no, word_no)
                        raw_words = None
                        try:
                            raw_words = page.get_text("words", sort=True)
                        except TypeError:
                            raw_words = page.get_text("words")
                        words_out: List[Dict[str, Any]] = []
                        for w in raw_words or []:
                            try:
                                x0, y0, x1, y1, wt = w[0], w[1], w[2], w[3], w[4]
                                s = (str(wt) if wt is not None else "").strip()
                                if not s:
                                    continue
                                words_out.append(
                                    {
                                        "x0": float(x0),
                                        "y0": float(y0),
                                        "x1": float(x1),
                                        "y1": float(y1),
                                        "text": s,
                                    }
                                )
                            except Exception as e:
                                logger.debug("Word tuple parse failed: %s", e)
                                continue
                        if words_out:
                            page_obj["words"] = words_out
                    except Exception as e:
                        logger.debug("Word layout extraction failed: %s", e)

                # Attach cropped rendering for vision-based UPR extraction (first N pages only).
                if upr_vision_enabled and int(page_num) <= int(upr_vision_max_pages):
                    try:
                        import base64

                        rect = page.rect
                        clip = fitz.Rect(
                            float(rect.x0),
                            float(rect.y0),
                            float(rect.x1),
                            float(rect.y0 + (rect.height * float(upr_vision_clip_top_frac))),
                        )
                        pix = page.get_pixmap(dpi=int(upr_vision_dpi), alpha=False, clip=clip)
                        png_bytes = pix.tobytes("png")
                        page_obj["upr_kpi_clip_png_b64"] = base64.b64encode(png_bytes).decode("ascii")
                        page_obj["upr_kpi_clip_box"] = [float(clip.x0), float(clip.y0), float(clip.x1), float(clip.y1)]
                    except Exception as e:
                        logger.debug("UPR KPI clip extraction failed: %s", e)

                result['pages'].append(page_obj)

                # Extract images if requested
                if extract_images:
                    images = self._extract_images_from_page(page, page_num)
                    result['images'].extend(images)

            # Record how many pages were actually processed (useful when max_pages is applied).
            try:
                if isinstance(result.get("metadata"), dict):
                    result["metadata"]["processed_pages"] = len(result.get("pages") or [])
            except Exception as e:
                logger.debug("Metadata processed_pages update failed: %s", e)

            # Extract table of contents / sections
            toc = doc.get_toc()
            if toc:
                result['sections'] = [
                    {
                        'level': level,
                        'title': title,
                        'page_number': page_num
                    }
                    for level, title, page_num in toc
                ]

            doc.close()

        except Exception as e:
            if self._is_probably_corrupt_pdf_error(e):
                raise DocumentProcessingError(
                    "PDF processing error: source PDF appears corrupted or structurally invalid."
                )
            raise DocumentProcessingError("PDF processing error.")

        return result

    def _extract_text_excluding_bboxes(self, page, exclude_bboxes: List[List[float]]) -> str:
        """
        Extract readable page text but drop blocks that intersect any excluded bbox.

        This is primarily used to avoid duplicating table content in both:
        - semantic text chunks (often messy for tables)
        - structured table chunks (cleaner and preferred)
        """
        try:
            blocks = None
            try:
                blocks = page.get_text("blocks", sort=True)
            except TypeError:
                blocks = page.get_text("blocks")
        except Exception as e:
            logger.debug("blocks extraction failed: %s", e)
            try:
                return page.get_text("text")
            except Exception as e2:
                logger.debug("page.get_text fallback failed: %s", e2)
                return ""

        norm_boxes: List[List[float]] = []
        for b in exclude_bboxes or []:
            try:
                x0, y0, x1, y1 = [float(v) for v in b]
                norm_boxes.append([x0, y0, x1, y1])
            except Exception as e:
                logger.debug("bbox float parse failed: %s", e)
                continue

        def _intersects(a: List[float], b: List[float], pad: float = 1.0) -> bool:
            ax0, ay0, ax1, ay1 = a
            bx0, by0, bx1, by1 = b
            # Expand excluded bbox slightly to remove bordering text
            bx0 -= pad
            by0 -= pad
            bx1 += pad
            by1 += pad
            return not (ax1 <= bx0 or ax0 >= bx1 or ay1 <= by0 or ay0 >= by1)

        kept: List[tuple] = []
        for blk in blocks or []:
            try:
                x0, y0, x1, y1, txt = blk[0], blk[1], blk[2], blk[3], blk[4]
            except Exception as e:
                logger.debug("block tuple parse failed: %s", e)
                continue
            if not txt or not str(txt).strip():
                continue
            bb = [float(x0), float(y0), float(x1), float(y1)]
            if any(_intersects(bb, tb) for tb in norm_boxes):
                continue
            kept.append((bb[1], bb[0], str(txt).strip()))  # sort by y then x

        kept.sort(key=lambda t: (t[0], t[1]))
        out = "\n".join(t[2] for t in kept if t[2]).strip()
        return out

    def _extract_tables_from_page(self, page, page_num: int) -> List[Dict[str, Any]]:
        """
        Best-effort table extraction for a PDF page.

        Uses PyMuPDF's table finder when available (version-dependent).
        Returns a list of tables, each with rows and optional bbox.
        """
        tables_out: List[Dict[str, Any]] = []

        def _is_spurious_label_table(*, rows_raw: List[List[str]], bbox: Optional[Any]) -> bool:
            """
            PyMuPDF's table finder can misclassify single "label boxes" (e.g. a framed 'Total 51.6M CHF')
            inside infographic-style visuals as a table, producing low-quality rows like 'Total 5'.

            These are not real tables and should be ignored so downstream consumers rely on
            template-specific visual extraction (e.g. UPR funding requirements) instead.
            """
            try:
                # Basic shape guardrail: only consider very small grids
                non_empty_rows = [r for r in (rows_raw or []) if isinstance(r, list) and any((c or "").strip() for c in r)]
                if not non_empty_rows:
                    return False
                nrows = len(non_empty_rows)
                ncols = max((len(r) for r in non_empty_rows), default=0)
                if nrows > 2 or ncols > 3:
                    return False

                flat = " ".join((c or "").strip() for r in non_empty_rows for c in r if (c or "").strip())
                low = flat.lower()
                if "total" not in low:
                    return False

                # Remove common noise tokens and amounts; if nothing meaningful remains, it's a label.
                scrub = low
                scrub = re.sub(r"\b(total|chf)\b", " ", scrub, flags=re.IGNORECASE)
                scrub = re.sub(r"\b\d+(?:\.\d+)?\s*[mk]\b", " ", scrub, flags=re.IGNORECASE)
                scrub = re.sub(r"\b\d{1,3}(?:,\d{3})+\b", " ", scrub)
                scrub = re.sub(r"[^a-z]+", " ", scrub)
                scrub = re.sub(r"\s+", " ", scrub).strip()
                if scrub:
                    # If there are still meaningful words (e.g. real headers), keep as a table.
                    if any(w in scrub for w in ("year", "national", "society", "confirmed", "requirement", "funding")):
                        return False

                # If bbox is available and small, it's almost certainly a label box.
                try:
                    if bbox is not None and len(bbox) == 4:
                        x0, y0, x1, y1 = [float(v) for v in bbox]
                        w = x1 - x0
                        h = y1 - y0
                        if (w > 0 and h > 0) and (w < 280.0 and h < 140.0):
                            return True
                except Exception as e:
                    logger.debug("Label box bbox check failed: %s", e)

                # Without bbox, fall back to the "scrubbed empty" heuristic.
                return (scrub == "")
            except Exception as e:
                logger.debug("_is_spurious_label_table failed: %s", e)
                return False

        # PyMuPDF table API is not available in all versions.
        find_tables = getattr(page, "find_tables", None)
        if not callable(find_tables):
            return tables_out

        finder = find_tables()
        # API shape differs slightly across versions; handle defensively.
        raw_tables = getattr(finder, "tables", None) or []
        for idx, t in enumerate(raw_tables):
            try:
                rows = None
                if hasattr(t, "extract"):
                    rows = t.extract()
                if not rows:
                    continue

                bbox = getattr(t, "bbox", None)
                # Normalize rows to list[list[str]]
                norm_rows: List[List[str]] = []
                for r in rows:
                    if r is None:
                        continue
                    if isinstance(r, (list, tuple)):
                        norm_rows.append([("" if c is None else str(c)).strip() for c in r])
                    else:
                        norm_rows.append([str(r).strip()])

                if not norm_rows:
                    continue

                header, rows_display, rows_expanded, records = self._normalize_table_grid(norm_rows)

                cell_fills = None
                if current_app.config.get("AI_TABLE_EXTRACT_COLORS_ENABLED", False):
                    try:
                        # Use the table bbox and a uniform grid approximation for fills.
                        bbox_list = list(bbox) if bbox is not None else None
                        # Note: colors are aligned to the display grid, not expanded rows.
                        if bbox_list and header and rows_display:
                            cell_fills = self._extract_table_cell_fills(
                                page=page,
                                table_bbox=bbox_list,
                                nrows=1 + len(rows_display),
                                ncols=len(header),
                            )
                    except Exception as e_colors:
                        logger.debug(f"Cell color extraction failed for page {page_num}: {e_colors}", exc_info=True)

                bbox_list = list(bbox) if bbox is not None else None
                if _is_spurious_label_table(rows_raw=norm_rows, bbox=bbox_list):
                    # Skip noisy/incorrect "tables" that are actually infographic labels.
                    continue

                tables_out.append(
                    {
                        "page_number": int(page_num),
                        "table_index": int(idx),
                        "bbox": bbox_list,
                        # Keep raw rows for debugging
                        "rows_raw": norm_rows,
                        # Collapsed/merged display rows (closest to extractor output)
                        "rows_display": rows_display,
                        # Expanded rows (per-year when possible) - preferred for chunking/embeddings
                        "header": header,
                        "rows": rows_expanded,
                        # Parsed record objects for deterministic consumption
                        "records": records,
                        "cell_fills": cell_fills,
                        "row_count": len(rows_expanded),
                        "col_count": max((len(r) for r in ([header] + rows_expanded)), default=0),
                    }
                )
            except Exception as e:
                logger.debug("Table extraction iteration failed: %s", e)
                continue

        return tables_out

    def _extract_table_cell_fills(self, *, page, table_bbox: List[float], nrows: int, ncols: int) -> List[List[Optional[str]]]:
        """
        Approximate cell background colors by raster sampling.

        Returns a matrix [nrows][ncols] with hex colors (e.g., "#FFA500") or None.
        """
        try:
            from PIL import Image
            import io
        except Exception as e:
            logger.debug("PIL import failed for cell fills: %s", e)
            return [[None for _ in range(ncols)] for _ in range(nrows)]

        # Render at moderate DPI for reasonable sampling quality.
        pix = page.get_pixmap(dpi=150, alpha=False)
        img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")

        page_rect = page.rect
        sx = img.width / float(page_rect.width or 1.0)
        sy = img.height / float(page_rect.height or 1.0)

        x0, y0, x1, y1 = [float(v) for v in table_bbox]
        if x1 <= x0 or y1 <= y0 or nrows <= 0 or ncols <= 0:
            return [[None for _ in range(ncols)] for _ in range(nrows)]

        cell_w = (x1 - x0) / float(ncols)
        cell_h = (y1 - y0) / float(nrows)

        def to_hex(rgb):
            try:
                r, g, b = rgb
                return f"#{int(r):02X}{int(g):02X}{int(b):02X}"
            except Exception as e:
                logger.debug("to_hex failed: %s", e)
                return None

        def sample_cell(cx0, cy0, cx1, cy1):
            # Sample a few points away from the center to avoid text.
            pts = [
                (0.15, 0.15),
                (0.85, 0.15),
                (0.15, 0.85),
                (0.85, 0.85),
                (0.50, 0.20),
            ]
            colors = []
            for pxr, pyr in pts:
                px = cx0 + (cx1 - cx0) * pxr
                py = cy0 + (cy1 - cy0) * pyr
                ix = max(0, min(img.width - 1, int(px * sx)))
                iy = max(0, min(img.height - 1, int(py * sy)))
                colors.append(img.getpixel((ix, iy)))
            # Mode (most frequent) color
            try:
                return max(set(colors), key=colors.count)
            except Exception as e:
                logger.debug("sample_cell mode failed: %s", e)
                return colors[0] if colors else (255, 255, 255)

        fills: List[List[Optional[str]]] = []
        for r in range(nrows):
            row_fills: List[Optional[str]] = []
            for c in range(ncols):
                cx0 = x0 + c * cell_w
                cy0 = y0 + r * cell_h
                cx1 = x0 + (c + 1) * cell_w
                cy1 = y0 + (r + 1) * cell_h
                rgb = sample_cell(cx0, cy0, cx1, cy1)
                hx = to_hex(rgb)
                row_fills.append(hx)
            fills.append(row_fills)
        return fills

    def _normalize_table_grid(
        self, grid: List[List[str]]
    ) -> Tuple[List[str], List[List[str]], List[List[str]], List[Dict[str, Any]]]:
        """
        Normalize a table grid into:
        - header: List[str]
        - rows_display: List[List[str]] (body rows only, empties removed; merged for readability)
        - rows_expanded: List[List[str]] (expanded per-year when possible; aligns with records)
        - records: List[Dict[str, Any]] parsed from the first column when possible

        This keeps the grid shape for transparency, but makes it far more usable for RAG.
        """
        if not grid:
            return [], [], [], []

        def norm_cell(v: Any) -> str:
            s = "" if v is None else str(v)
            s = re.sub(r"\s+", " ", s).strip()
            return s

        # Normalize cell whitespace and rectangularize row lengths
        max_cols = max((len(r) for r in grid if isinstance(r, list)), default=0)
        rows = []
        for r in grid:
            if not isinstance(r, list):
                continue
            rr = [norm_cell(c) for c in r]
            if len(rr) < max_cols:
                rr = rr + [""] * (max_cols - len(rr))
            rows.append(rr)

        # Drop fully empty rows
        rows = [r for r in rows if any(c for c in r)]
        if not rows:
            return [], [], [], []

        header = rows[0]
        body = rows[1:] if len(rows) > 1 else []

        # Header fixups
        # 1) Replace the known "collapsed" first header cell with a clearer label (still a single column).
        if header and header[0]:
            h0 = header[0].lower()
            if ("national society" in h0) and ("year" in h0) and ("fund" in h0):
                header[0] = "National Society / Year / Funding Requirement / Confirmed Funding"

        # 2) If the next row is a tiny "continuation" row (e.g., shows 'V' 'E' to complete words),
        # use it to patch headers and then drop it.
        if body:
            maybe = body[0]
            # Continuation row heuristic: very short cells and mostly empties
            short_cells = sum(1 for c in maybe if c and len(c) <= 2)
            filled = sum(1 for c in maybe if c)
            if filled > 0 and short_cells == filled and filled <= max(2, max_cols // 3):
                for i, c in enumerate(maybe):
                    if not c:
                        continue
                    # If header appears to be missing leading character (starts lowercase), prefix the continuation.
                    if i < len(header) and header[i] and header[i][:1].islower() and c.isalpha() and len(c) == 1:
                        header[i] = f"{c}{header[i]}".strip()
                # Drop the continuation row
                body = body[1:]

        # Group continuation rows (empty first column but other columns filled) under their anchor.
        anchors: List[Dict[str, Any]] = []
        current_anchor: Optional[Dict[str, Any]] = None

        for r in body:
            if not any(r):
                continue
            if r[0]:
                current_anchor = {"row": r[:], "continuations": []}
                anchors.append(current_anchor)
                continue
            if current_anchor is not None and any(c for c in r[1:]):
                current_anchor["continuations"].append(r[:])
                continue
            # Unattached noise row; ignore.

        # Build a readable merged grid for display/embedding (do NOT concatenate conflicting values).
        merged: List[List[str]] = []
        for a in anchors:
            base = (a.get("row") or [])[:]
            conts = a.get("continuations") or []
            for c in conts:
                for i in range(1, max_cols):
                    if c[i] and not base[i]:
                        base[i] = c[i]
                    # If both exist and differ, keep base as-is (avoid "a | b" artifacts).
            merged.append(base)

        # Parse records from merged rows
        col_names = header[1:] if len(header) > 1 else []
        records: List[Dict[str, Any]] = []
        rows_expanded: List[List[str]] = []

        for idx_row, r in enumerate(merged):
            left = r[0]
            if not left:
                continue

            years = [int(y) for y in _YEAR_RE.findall(left)]

            # Derive society name by stripping years and amounts (handles wrapped words like "... Red Cross")
            society = left
            try:
                society = _YEAR_RE.sub(" ", society)
                society = _AMOUNT_RE.sub(" ", society)
                society = re.sub(r"\s+", " ", society).strip()
            except Exception as e:
                logger.debug("society regex strip failed: %s", e)
                society = society.strip()

            # Category values from the merged display row (base values only)
            base_categories: Dict[str, Any] = {}
            for i, name in enumerate(col_names, start=1):
                val = r[i] if i < len(r) else ""
                if val:
                    base_categories[name] = val

            # Continuation rows for this anchor (if available): use them to split per-year categories
            cont_rows: List[List[str]] = []
            try:
                cont_rows = anchors[idx_row].get("continuations") or []
            except Exception as e:
                logger.debug("continuations get failed: %s", e)
                cont_rows = []

            cont_categories_list: List[Dict[str, Any]] = []
            for cr in cont_rows:
                cat: Dict[str, Any] = {}
                for i, name in enumerate(col_names, start=1):
                    val = cr[i] if i < len(cr) else ""
                    if val:
                        cat[name] = val
                if cat:
                    cont_categories_list.append(cat)

            if years:
                # Parse per-year amount segments from the left cell:
                # e.g. "2026 4.2M 4.2M 2027 4.2M 4.2M" => req+confirmed per year
                # or   "2026 1M 2027 1M" => requirement only per year
                year_matches = list(_YEAR_RE.finditer(left))
                year_to_amounts: Dict[int, List[str]] = {}
                for i, m in enumerate(year_matches):
                    y = int(m.group(1))
                    start = m.end()
                    end = year_matches[i + 1].start() if i + 1 < len(year_matches) else len(left)
                    seg = left[start:end]
                    seg_amounts = [a.strip() for a in _AMOUNT_RE.findall(seg)]
                    year_to_amounts[y] = seg_amounts

                for y_idx, y in enumerate(years):
                    rec: Dict[str, Any] = {"national_society": society, "year": y}
                    seg_amounts = year_to_amounts.get(int(y), [])
                    if seg_amounts:
                        # If two+ values appear for this year, interpret as requirement + confirmed.
                        # If only one appears, treat it as requirement (confirmed may be blank in table).
                        rec["funding_requirement"] = seg_amounts[0]
                        if len(seg_amounts) >= 2:
                            rec["confirmed_funding"] = seg_amounts[1]

                    # Assign categories:
                    # - If we have one continuation row, apply it to all years (common).
                    # - If we have >= N continuation rows for N years, map by year index (fixes the "a | b" issue).
                    # - Otherwise fall back to base_categories.
                    cats: Dict[str, Any] = {}
                    # If there are multiple years but we don't have per-year continuation rows,
                    # base categories likely belong to the first year only (common extraction artifact).
                    if base_categories and (len(years) == 1 or cont_categories_list or y_idx == 0):
                        cats.update(base_categories)
                    if cont_categories_list:
                        if len(cont_categories_list) >= len(years):
                            cats.update(cont_categories_list[y_idx])
                        elif len(cont_categories_list) == 1:
                            cats.update(cont_categories_list[0])
                        else:
                            # Best effort: map first K rows to first K years
                            if y_idx < len(cont_categories_list):
                                cats.update(cont_categories_list[y_idx])
                    if cats:
                        rec["categories"] = cats
                    records.append(rec)

                    # Expanded row aligned with this record
                    first_parts = [society, str(y)]
                    if rec.get("funding_requirement"):
                        first_parts.append(str(rec["funding_requirement"]))
                    if rec.get("confirmed_funding"):
                        first_parts.append(str(rec["confirmed_funding"]))
                    out_row = [" ".join(first_parts).strip()]
                    for name in col_names:
                        out_row.append(str(cats.get(name, "")) if cats and name in cats else "")
                    rows_expanded.append(out_row)
            else:
                rec = {"national_society": society}
                cats: Dict[str, Any] = {}
                if base_categories:
                    cats.update(base_categories)
                if cont_categories_list:
                    # Merge all continuation categories (best effort)
                    for c in cont_categories_list:
                        cats.update(c)
                if cats:
                    rec["categories"] = cats
                records.append(rec)
                # If we can't split by year, keep the merged row.
                out_row = r[:]
                rows_expanded.append(out_row)

        # Normalize expanded rows to match header width
        expected_cols = max(1, len(header))
        normalized_expanded: List[List[str]] = []
        for rr in rows_expanded:
            rr = (rr or [])[:expected_cols] + [""] * max(0, expected_cols - len(rr or []))
            normalized_expanded.append(rr)

        return header, merged, normalized_expanded, records

    def _process_word(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process Word document (.docx)."""
        try:
            from docx import Document
        except ImportError:
            raise DocumentProcessingError("python-docx not installed. Run: pip install python-docx")

        result = {
            'text': '',
            'metadata': {},
            'images': [],
            'sections': [],
            'pages': []  # Word doesn't have explicit pages
        }

        try:
            doc = Document(file_path)

            # Extract metadata
            core_properties = doc.core_properties
            result['metadata'] = {
                'title': core_properties.title or filename,
                'author': core_properties.author or '',
                'subject': core_properties.subject or '',
                'created': str(core_properties.created) if core_properties.created else '',
                'modified': str(core_properties.modified) if core_properties.modified else '',
                'total_paragraphs': len(doc.paragraphs),
            }

            # Extract text with section tracking
            current_section = None
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Detect headings (sections)
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name.replace('Heading ', '')) if para.style.name[-1].isdigit() else 1
                    result['sections'].append({
                        'level': level,
                        'title': text,
                        'paragraph_index': len(result['pages'])
                    })
                    current_section = text

                result['text'] += f"{text}\n\n"
                result['pages'].append({
                    'paragraph_index': len(result['pages']),
                    'text': text,
                    'section': current_section,
                    'style': para.style.name
                })

            # Extract images (simplified - full implementation would need more work)
            # for rel in doc.part.rels.values():
            #     if "image" in rel.target_ref:
            #         # Extract image data
            #         pass

        except Exception as e:
            raise DocumentProcessingError("Word processing error.")

        return result

    def _process_excel(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process Excel document."""
        try:
            import openpyxl
        except ImportError:
            raise DocumentProcessingError("openpyxl not installed. Run: pip install openpyxl")

        result = {
            'text': '',
            'metadata': {},
            'images': [],
            'sections': [],
            'pages': []  # Each sheet is a "page"
        }

        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)

            result['metadata'] = {
                'title': filename,
                'total_sheets': len(wb.sheetnames),
                'sheet_names': wb.sheetnames,
            }

            # Process each sheet
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]

                # Extract cell values as text
                sheet_text = f"\n\n=== Sheet: {sheet_name} ===\n\n"
                rows = []

                for row in ws.iter_rows(values_only=True):
                    # Filter out empty rows
                    if any(cell is not None for cell in row):
                        row_text = '\t'.join(str(cell) if cell is not None else '' for cell in row)
                        rows.append(row_text)
                        sheet_text += row_text + '\n'

                result['text'] += sheet_text
                result['pages'].append({
                    'sheet_name': sheet_name,
                    'text': sheet_text,
                    'row_count': len(rows)
                })

            wb.close()

        except Exception as e:
            raise DocumentProcessingError("Excel processing error.")

        return result

    def _process_text(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process plain text file."""
        result = {
            'text': '',
            'metadata': {},
            'images': [],
            'sections': [],
            'pages': []
        }

        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = None

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                raise DocumentProcessingError("Could not decode text file with any supported encoding")

            result['text'] = text
            result['metadata'] = {
                'title': filename,
                'char_count': len(text),
                'line_count': len(text.splitlines()),
            }

        except Exception as e:
            raise DocumentProcessingError("Text processing error.")

        return result

    def _process_markdown(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process Markdown file."""
        try:
            import markdown
            from bs4 import BeautifulSoup
        except ImportError:
            # Fallback to plain text processing
            return self._process_text(file_path, filename)

        result = {
            'text': '',
            'metadata': {},
            'images': [],
            'sections': [],
            'pages': []
        }

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                md_text = f.read()

            # Convert to HTML
            html = markdown.markdown(md_text, extensions=['extra', 'toc'])

            # Extract plain text
            soup = BeautifulSoup(html, 'html.parser')
            text = soup.get_text()

            # Extract headings as sections
            for heading in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                level = int(heading.name[1])
                result['sections'].append({
                    'level': level,
                    'title': heading.get_text().strip()
                })

            result['text'] = text
            result['metadata'] = {
                'title': result['sections'][0]['title'] if result['sections'] else filename,
                'char_count': len(text),
                'heading_count': len(result['sections']),
            }

        except Exception as e:
            raise DocumentProcessingError("Markdown processing error.")

        return result

    def _process_html(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process HTML file."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise DocumentProcessingError("beautifulsoup4 not installed. Run: pip install beautifulsoup4")

        result = {
            'text': '',
            'metadata': {},
            'images': [],
            'sections': [],
            'pages': []
        }

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html = f.read()

            soup = BeautifulSoup(html, 'html.parser')

            # Extract metadata from HTML head
            title = soup.find('title')
            result['metadata'] = {
                'title': title.get_text() if title else filename,
            }

            # Extract text
            text = soup.get_text(separator='\n\n')
            result['text'] = text

            # Extract headings as sections
            for heading in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                level = int(heading.name[1])
                result['sections'].append({
                    'level': level,
                    'title': heading.get_text().strip()
                })

        except Exception as e:
            raise DocumentProcessingError("HTML processing error.")

        return result

    def _ocr_page(self, page, page_num: int) -> str:
        """Perform OCR on a PDF page (for scanned documents)."""
        if not current_app.config.get('AI_OCR_ENABLED', False):
            return ""

        try:
            import pytesseract
            from PIL import Image
            import io

            # Render page as image
            pix = page.get_pixmap(dpi=300)
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))

            # Perform OCR
            text = pytesseract.image_to_string(img)
            logger.info(f"OCR performed on page {page_num}, extracted {len(text)} characters")
            return text

        except ImportError:
            logger.warning("pytesseract not installed, skipping OCR")
            return ""
        except Exception as e:
            logger.warning(f"OCR failed on page {page_num}: {str(e)}")
            return ""

    def _extract_images_from_page(self, page, page_num: int) -> List[Dict[str, Any]]:
        """Extract images from a PDF page."""
        images = []

        try:
            image_list = page.get_images(full=True)

            for img_index, img in enumerate(image_list):
                xref = img[0]
                images.append({
                    'page_number': page_num,
                    'image_index': img_index,
                    'xref': xref,
                    # Image data would be extracted here if needed
                })

        except Exception as e:
            logger.warning(f"Failed to extract images from page {page_num}: {str(e)}")

        return images
