from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def is_table_line(line: str) -> bool:
    return "|" in line and line.strip().startswith("|")


def parse_table_row(line: str) -> list[str]:
    raw = line.strip().strip("|")
    return [cell.strip() for cell in raw.split("|")]


def is_table_separator(line: str) -> bool:
    # Example: |---|---| or |:---|---:|
    cleaned = line.strip().replace(" ", "")
    return bool(re.fullmatch(r"\|[:\-|]+\|?", cleaned))


def apply_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)


def add_title_line(doc: Document, text: str, size: int) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.strip())
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(size)


def add_paragraph_with_style(doc: Document, text: str, in_toc: bool = False) -> None:
    stripped = text.rstrip()
    if not stripped:
        doc.add_paragraph("")
        return

    if in_toc:
        doc.add_paragraph(stripped.strip())
        return

    heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
    if heading_match:
        heading_text = heading_match.group(2).strip()
        level = min(len(heading_match.group(1)), 4)
        if re.match(r"^\d+\.\s+", heading_text):
            doc.add_heading(heading_text, level=1)
        elif re.match(r"^\d+\.\d+", heading_text):
            doc.add_heading(heading_text, level=2)
        else:
            doc.add_heading(heading_text, level=level)
        return

    if re.match(r"^\d+\.\s+", stripped) and not stripped.endswith(":"):
        doc.add_paragraph(stripped, style="List Number")
        return

    if stripped.startswith("- "):
        doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        return

    if stripped.startswith("> "):
        doc.add_paragraph(stripped[2:].strip(), style="Intense Quote")
        return

    if re.fullmatch(r"^[-_*]{3,}$", stripped):
        doc.add_paragraph("")
        return

    doc.add_paragraph(stripped)


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    apply_base_styles(doc)

    # Kobo-like title page rendering for the first two H1 lines.
    title_h1_lines = [idx for idx, line in enumerate(lines) if re.match(r"^#\s+", line)]
    if len(title_h1_lines) >= 2 and title_h1_lines[0] == 0 and title_h1_lines[1] == 1:
        add_title_line(doc, re.sub(r"^#\s+", "", lines[0]), size=18)
        add_title_line(doc, re.sub(r"^#\s+", "", lines[1]), size=15)
        doc.add_paragraph("")
        i = 2
    else:
        i = 0

    in_toc = False
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "## Table of Contents":
            doc.add_heading("Table of Contents", level=2)
            in_toc = True
            i += 1
            continue
        if in_toc and re.fullmatch(r"^[-_*]{3,}$", stripped):
            in_toc = False
            doc.add_page_break()
            i += 1
            continue

        if is_table_line(line):
            table_lines: list[str] = []
            while i < len(lines) and is_table_line(lines[i]):
                table_lines.append(lines[i])
                i += 1

            rows = [parse_table_row(tl) for tl in table_lines if not is_table_separator(tl)]
            if rows:
                col_count = max(len(r) for r in rows)
                table = doc.add_table(rows=0, cols=col_count)
                table.style = "Table Grid"
                for row_idx, row_data in enumerate(rows):
                    row = table.add_row().cells
                    for col_idx in range(col_count):
                        row[col_idx].text = row_data[col_idx] if col_idx < len(row_data) else ""
                    if row_idx == 0:
                        for cell in row:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.bold = True
            continue

        add_paragraph_with_style(doc, line, in_toc=in_toc)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: python scripts/md_to_docx.py <input.md> <output.docx>")
        return 1

    md_path = Path(sys.argv[1])
    docx_path = Path(sys.argv[2])
    if not md_path.exists():
        print(f"Input file not found: {md_path}")
        return 1

    md_to_docx(md_path, docx_path)
    print(f"Created: {docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
